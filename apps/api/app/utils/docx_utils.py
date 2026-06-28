from io import BytesIO
from docx import Document


def create_docx_from_text(content: str, title: str) -> bytes:
    """Create a DOCX document from plain text.

    Args:
        content: Document text content.
        title: Document title.

    Returns:
        DOCX bytes.
    """
    doc = Document()
    doc.add_heading(title, level=0)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
